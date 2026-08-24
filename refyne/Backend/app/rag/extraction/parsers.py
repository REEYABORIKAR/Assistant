import docx
import pandas as pd
from pypdf import PdfReader


class ExtractionError(Exception):
    pass

def extract_pdf(file_path: str) -> list[dict]:
    results = []
    try:
        reader = PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                results.append({
                    "text": text,
                    "metadata": {"page_number": i + 1}
                })
    except Exception as e:
        raise ExtractionError(f"PDF extraction failed: {str(e)}")
    return results

def extract_docx(file_path: str) -> list[dict]:
    results = []
    try:
        doc = docx.Document(file_path)
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                results.append({
                    "text": text,
                    "metadata": {"paragraph_index": i}
                })
        # Extract tables as well
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    results.append({
                        "text": " | ".join(row_data),
                        "metadata": {"table_index": i, "row_index": j}
                    })
    except Exception as e:
        raise ExtractionError(f"DOCX extraction failed: {str(e)}")
    return results

def extract_txt(file_path: str) -> list[dict]:
    try:
        with open(file_path, encoding="utf-8") as f:
            text = f.read()
            return [{"text": text, "metadata": {}}]
    except UnicodeDecodeError:
        try:
            with open(file_path, encoding="latin-1") as f:
                text = f.read()
                return [{"text": text, "metadata": {}}]
        except Exception as e:
            raise ExtractionError(f"TXT extraction failed: {str(e)}")
    except Exception as e:
        raise ExtractionError(f"TXT extraction failed: {str(e)}")

def extract_csv(file_path: str) -> list[dict]:
    results = []
    try:
        df = pd.read_csv(file_path)
        for index, row in df.iterrows():
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
            if row_text:
                results.append({
                    "text": row_text,
                    "metadata": {"row_index": index}
                })
    except Exception as e:
        raise ExtractionError(f"CSV extraction failed: {str(e)}")
    return results

def extract_xlsx(file_path: str) -> list[dict]:
    results = []
    try:
        xls = pd.ExcelFile(file_path)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            for index, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                if row_text:
                    results.append({
                        "text": row_text,
                        "metadata": {"sheet_name": sheet_name, "row_index": index}
                    })
    except Exception as e:
        raise ExtractionError(f"XLSX extraction failed: {str(e)}")
    return results

def extract_doc(file_path: str) -> list[dict]:
    # Best effort for legacy .doc
    raise ExtractionError("Legacy .doc extraction is unavailable in the current environment. Please convert the document to .docx.")

def extract_document(file_path: str, ext: str) -> list[dict]:
    ext = ext.lower()
    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext == ".docx":
        return extract_docx(file_path)
    elif ext == ".txt":
        return extract_txt(file_path)
    elif ext == ".csv":
        return extract_csv(file_path)
    elif ext == ".xlsx":
        return extract_xlsx(file_path)
    elif ext == ".doc":
        return extract_doc(file_path)
    else:
        raise ExtractionError(f"Unsupported extension: {ext}")
